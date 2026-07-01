"""
Extracts plain text from uploaded BRD/FSD documents (PDF or DOCX).
"""
import os
import pdfplumber
from docx import Document


def extract_text(file_path: str) -> str:
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return _extract_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return _extract_docx(file_path)
    elif ext in (".txt", ".md"):
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    else:
        raise ValueError(f"Unsupported document type: {ext}")


def _extract_pdf(file_path: str) -> str:
    text_parts = []
    with pdfplumber.open(file_path) as pdf:
        for i, page in enumerate(pdf.pages):
            page_text = page.extract_text() or ""
            text_parts.append(f"--- Page {i + 1} ---\n{page_text}")
    return "\n\n".join(text_parts)


def _extract_docx(file_path: str) -> str:
    doc = Document(file_path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            parts.append(" | ".join(cells))
    return "\n".join(parts)


def truncate_for_llm(text: str, max_chars: int = 18000) -> str:
    """Keep documents from blowing the context window. Keeps head + tail."""
    if len(text) <= max_chars:
        return text
    half = max_chars // 2
    return (
        text[:half]
        + "\n\n... [TRUNCATED MIDDLE — document continues] ...\n\n"
        + text[-half:]
    )
